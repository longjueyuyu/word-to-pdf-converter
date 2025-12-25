"""创建测试Word文档"""
from docx import Document
import os

# 创建测试文档目录
test_dir = "test_docs"
os.makedirs(test_dir, exist_ok=True)

# 创建3个测试Word文档
for i in range(1, 4):
    doc = Document()
    doc.add_heading(f'测试文档 {i}', 0)
    doc.add_paragraph(f'这是第 {i} 个测试文档的内容。')
    doc.add_paragraph('这个文档用于测试Word转PDF功能。')
    
    # 添加一些额外内容
    doc.add_heading('功能说明', level=1)
    doc.add_paragraph('• 支持批量转换Word文件为PDF')
    doc.add_paragraph('• 显示转换进度')
    doc.add_paragraph('• 提供详细日志')
    
    filename = os.path.join(test_dir, f'测试文档_{i}.docx')
    doc.save(filename)
    print(f"✓ 创建: {filename}")

print(f"\n成功创建了 3 个测试Word文档！")
